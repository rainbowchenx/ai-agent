"""知识库服务类，整合关系数据库和向量数据库。

该模块提供了知识库的完整服务，包括文档上传、处理、存储、检索等功能。
"""

import os
import re
import uuid
import shutil
from typing import List, Optional, Dict, Any
from pathlib import Path
from datetime import datetime, UTC

from fastapi import HTTPException, UploadFile
from sqlmodel import Session, select
import PyPDF2
from docx import Document
import markdown

from app.core.config import settings
from app.core.logging import logger
from app.models.knowledge import KnowledgeDocument, KnowledgeChunk
from app.services.database import database_service
from app.services.vector_store import vector_store_service


class KnowledgeService:
    """知识库服务类。
    
    提供知识库的完整功能，包括文档管理、文本处理、向量化等。
    """

    def __init__(self):
        """初始化知识库服务。"""
        # 创建文档存储目录
        self.upload_dir = Path(settings.CHROMA_PERSIST_DIRECTORY) / "uploads"
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def upload_document(
        self, 
        file: UploadFile, 
        description: str, 
        user_id: int
    ) -> KnowledgeDocument:
        """上传并处理文档。
        
        Args:
            file: 上传的文件
            description: 文档描述
            user_id: 上传用户ID
            
        Returns:
            KnowledgeDocument: 创建的文档记录
            
        Raises:
            HTTPException: 当文件处理失败时
        """
        try:
            # 生成文档ID
            document_id = str(uuid.uuid4())
            
            # 保存文件
            file_path = self.upload_dir / f"{document_id}_{file.filename}"
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            # 创建文档记录
            document = KnowledgeDocument(
                id=document_id,
                name=file.filename,
                description=description,
                file_type=file.content_type or self._get_file_type(file.filename),
                file_size=file.size,
                file_path=str(file_path),
                user_id=user_id,
                status="processing"
            )
            
            # 保存到数据库
            with Session(database_service.engine) as session:
                session.add(document)
                session.commit()
                session.refresh(document)
            
            logger.info(
                "document_uploaded",
                document_id=document_id,
                filename=file.filename,
                user_id=user_id
            )
            
            # 异步处理文档（在实际应用中应该使用后台任务）
            await self._process_document(document)
            
            return document
            
        except Exception as e:
            logger.error(
                "document_upload_error",
                filename=file.filename,
                user_id=user_id,
                error=str(e)
            )
            raise HTTPException(status_code=500, detail="文档上传失败")

    async def _process_document(self, document: KnowledgeDocument) -> bool:
        """处理文档，提取文本并向量化。
        
        Args:
            document: 要处理的文档
            
        Returns:
            bool: 处理成功返回True，否则返回False
        """
        try:
            # 提取文本
            text_content = await self._extract_text(document.file_path, document.file_type)
            if not text_content:
                raise Exception("无法提取文档文本内容")
            
            # 分块处理
            chunks = await self._create_chunks(document.id, text_content, document.file_type)
            if not chunks:
                raise Exception("文档分块失败")
            
            # 保存分块到数据库
            with Session(database_service.engine) as session:
                for chunk in chunks:
                    session.add(chunk)
                session.commit()
            
            # 添加到向量数据库
            success = await vector_store_service.add_document_chunks(document.id, chunks)
            if not success:
                raise Exception("向量化失败")
            
            # 更新文档状态
            with Session(database_service.engine) as session:
                doc = session.get(KnowledgeDocument, document.id)
                if doc:
                    doc.status = "completed"
                    doc.chunks = len(chunks)
                    doc.vector_count = len(chunks)
                    session.add(doc)
                    session.commit()
            
            logger.info(
                "document_processed",
                document_id=document.id,
                chunk_count=len(chunks)
            )
            return True
            
        except Exception as e:
            logger.error(
                "document_processing_error",
                document_id=document.id,
                error=str(e)
            )
            
            # 更新文档状态为失败
            with Session(database_service.engine) as session:
                doc = session.get(KnowledgeDocument, document.id)
                if doc:
                    doc.status = "failed"
                    doc.error_message = str(e)
                    session.add(doc)
                    session.commit()
            
            return False

    async def _extract_text(self, file_path: str, file_type: str) -> Optional[str]:
        """从文件中提取文本内容。
        
        Args:
            file_path: 文件路径
            file_type: 文件类型
            
        Returns:
            Optional[str]: 提取的文本内容
        """
        try:
            if file_type == "application/pdf":
                return self._extract_pdf_text(file_path)
            elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                return self._extract_docx_text(file_path)
            elif file_type == "text/plain":
                return self._extract_txt_text(file_path)
            elif file_type == "text/markdown":
                return self._extract_markdown_text(file_path)
            else:
                logger.warning("unsupported_file_type", file_type=file_type)
                return None
                
        except Exception as e:
            logger.error("text_extraction_error", file_path=file_path, error=str(e))
            return None

    def _extract_pdf_text(self, file_path: str) -> str:
        """提取PDF文件文本。"""
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text

    def _extract_docx_text(self, file_path: str) -> str:
        """提取DOCX文件文本。"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def _extract_txt_text(self, file_path: str) -> str:
        """提取TXT文件文本。"""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()

    def _extract_markdown_text(self, file_path: str) -> str:
        """提取Markdown文件文本。"""
        with open(file_path, "r", encoding="utf-8") as file:
            md_content = file.read()
            # 转换为纯文本
            html = markdown.markdown(md_content)
            # 简单的HTML标签移除
            text = re.sub(r'<[^>]+>', '', html)
            return text

    async def _create_chunks(
        self, 
        document_id: str, 
        text: str, 
        file_type: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> List[KnowledgeChunk]:
        """将文本分割成块。
        
        Args:
            document_id: 文档ID
            text: 文本内容
            file_type: 文件类型
            chunk_size: 块大小
            chunk_overlap: 块重叠大小
            
        Returns:
            List[KnowledgeChunk]: 文本块列表
        """
        chunks = []
        
        # 简单的文本分割（在实际应用中可以使用更复杂的分割策略）
        words = text.split()
        current_chunk = []
        current_size = 0
        chunk_index = 0
        
        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1  # +1 for space
            
            if current_size >= chunk_size:
                chunk_text = " ".join(current_chunk)
                chunk = KnowledgeChunk(
                    id=str(uuid.uuid4()),
                    document_id=document_id,
                    content=chunk_text,
                    chunk_index=chunk_index,
                    metadata={
                        "file_type": file_type,
                        "chunk_size": len(chunk_text),
                        "created_at": datetime.now(UTC).isoformat()
                    }
                )
                chunks.append(chunk)
                
                # 准备下一个块（包含重叠）
                overlap_words = current_chunk[-chunk_overlap//10:] if chunk_overlap > 0 else []
                current_chunk = overlap_words
                current_size = sum(len(word) + 1 for word in overlap_words)
                chunk_index += 1
        
        # 处理最后一个块
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk = KnowledgeChunk(
                id=str(uuid.uuid4()),
                document_id=document_id,
                content=chunk_text,
                chunk_index=chunk_index,
                metadata={
                    "file_type": file_type,
                    "chunk_size": len(chunk_text),
                    "created_at": datetime.now(UTC).isoformat()
                }
            )
            chunks.append(chunk)
        
        return chunks

    def _get_file_type(self, filename: str) -> str:
        """根据文件名获取文件类型。"""
        ext = Path(filename).suffix.lower()
        type_map = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain",
            ".md": "text/markdown"
        }
        return type_map.get(ext, "application/octet-stream")

    async def get_document_list(self, user_id: int) -> List[KnowledgeDocument]:
        """获取用户的文档列表。
        
        Args:
            user_id: 用户ID
            
        Returns:
            List[KnowledgeDocument]: 文档列表
        """
        with Session(database_service.engine) as session:
            statement = select(KnowledgeDocument).where(
                KnowledgeDocument.user_id == user_id
            ).order_by(KnowledgeDocument.created_at.desc())
            documents = session.exec(statement).all()
            return documents

    async def get_document(self, document_id: str, user_id: int) -> Optional[KnowledgeDocument]:
        """获取文档详情。
        
        Args:
            document_id: 文档ID
            user_id: 用户ID
            
        Returns:
            Optional[KnowledgeDocument]: 文档信息
        """
        with Session(database_service.engine) as session:
            statement = select(KnowledgeDocument).where(
                KnowledgeDocument.id == document_id,
                KnowledgeDocument.user_id == user_id
            )
            document = session.exec(statement).first()
            return document

    async def delete_document(self, document_id: str, user_id: int) -> bool:
        """删除文档。
        
        Args:
            document_id: 文档ID
            user_id: 用户ID
            
        Returns:
            bool: 删除成功返回True，否则返回False
        """
        try:
            # 获取文档信息
            document = await self.get_document(document_id, user_id)
            if not document:
                return False
            
            # 删除向量数据库中的数据
            await vector_store_service.delete_document_chunks(document_id)
            
            # 删除数据库记录
            with Session(database_service.engine) as session:
                # 删除分块
                statement = select(KnowledgeChunk).where(
                    KnowledgeChunk.document_id == document_id
                )
                chunks = session.exec(statement).all()
                for chunk in chunks:
                    session.delete(chunk)
                
                # 删除文档
                session.delete(document)
                session.commit()
            
            # 删除文件
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            logger.info("document_deleted", document_id=document_id, user_id=user_id)
            return True
            
        except Exception as e:
            logger.error(
                "document_delete_error",
                document_id=document_id,
                user_id=user_id,
                error=str(e)
            )
            return False

    async def search_knowledge(
        self, 
        query: str, 
        user_id: int,
        n_results: int = 5
    ) -> List[Dict[str, Any]]:
        """搜索知识库。
        
        Args:
            query: 搜索查询
            user_id: 用户ID
            n_results: 返回结果数量
            
        Returns:
            List[Dict[str, Any]]: 搜索结果
        """
        try:
            # 获取用户的文档ID列表
            documents = await self.get_document_list(user_id)
            document_ids = [doc.id for doc in documents if doc.status == "completed"]
            
            if not document_ids:
                return []
            
            # 搜索向量数据库
            results = await vector_store_service.search_similar_chunks(
                query, n_results, document_ids
            )
            
            # 添加文档信息
            for result in results:
                doc_id = result["metadata"]["document_id"]
                document = next((doc for doc in documents if doc.id == doc_id), None)
                if document:
                    result["document"] = {
                        "id": document.id,
                        "name": document.name,
                        "description": document.description
                    }
            
            return results
            
        except Exception as e:
            logger.error("knowledge_search_error", query=query, user_id=user_id, error=str(e))
            return []


# 创建一个单例实例
knowledge_service = KnowledgeService() 